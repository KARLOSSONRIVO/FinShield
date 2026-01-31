"""
Tesseract OCR Extractor - Text and layout extraction from documents.

Provides two main functions:
- extract_text_simple: Basic text extraction (for OCR and precheck)
- extract_text_with_layout: Full extraction with bounding boxes (for templates)
"""
import os
import tempfile
import shutil
from typing import Optional, List, Dict, Any
from pdf2image import convert_from_path
from docx import Document
import pytesseract
from PIL import Image
import pdfplumber


def _find_poppler_path() -> Optional[str]:
    """Find Poppler installation path on Windows."""
    if shutil.which("pdftoppm"):
        return None
    
    possible_paths = [
        r"C:\ProgramData\poppler\Library\bin",
        r"C:\Program Files\poppler\Library\bin",
        r"C:\Program Files\poppler\bin",
        r"C:\poppler\Library\bin",
        r"C:\poppler\bin",
        r"C:\tools\poppler\Library\bin",
        os.path.expanduser(r"~\poppler\Library\bin"),
    ]
    
    for path in possible_paths:
        if os.path.exists(os.path.join(path, "pdftoppm.exe")):
            return path
    
    return None


POPPLER_PATH = _find_poppler_path()


# =============================================================================
# SIMPLE TEXT EXTRACTION (for OCR service and precheck)
# =============================================================================

def extract_text_from_pdf(path: str) -> str:
    """Extract text from PDF using pdfplumber with Tesseract fallback."""
    text_blocks = []

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_blocks.append(text)
            else:
                # Fallback to OCR for scanned pages
                img = page.to_image(resolution=300).original
                ocr_text = pytesseract.image_to_string(img, config="--psm 6")
                if ocr_text:
                    text_blocks.append(ocr_text)

    return "\n".join(text_blocks)


def extract_text_from_docx(path: str) -> str:
    """Extract text from DOCX."""
    doc = Document(path)
    blocks = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    blocks.append(cell.text)

    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text)

    return "\n".join(blocks)


def extract_text_simple(path: str, filename: str) -> dict:
    """Simple text extraction for OCR and precheck services."""
    name = filename.lower()

    if name.endswith(".pdf"):
        return {
            "text": extract_text_from_pdf(path),
            "source": "tesseract"
        }

    if name.endswith(".docx"):
        return {
            "text": extract_text_from_docx(path),
            "source": "python-docx"
        }

    raise ValueError("UNSUPPORTED_FILE_TYPE")


# =============================================================================
# LAYOUT EXTRACTION (for template service)
# =============================================================================

def _normalize_bbox(bbox: Dict, img_width: int, img_height: int) -> Dict[str, float]:
    """Convert Tesseract bbox to normalized coordinates (0-1)."""
    x1 = bbox["left"] / img_width
    y1 = bbox["top"] / img_height
    x2 = (bbox["left"] + bbox["width"]) / img_width
    y2 = (bbox["top"] + bbox["height"]) / img_height
    
    return {
        "x1": round(x1, 4),
        "y1": round(y1, 4),
        "x2": round(x2, 4),
        "y2": round(y2, 4),
    }


def _get_position_label(bbox: Dict[str, float]) -> str:
    """Determine position label based on normalized bbox center."""
    center_x = (bbox["x1"] + bbox["x2"]) / 2
    center_y = (bbox["y1"] + bbox["y2"]) / 2
    
    if center_x < 0.33:
        h_pos = "left"
    elif center_x > 0.66:
        h_pos = "right"
    else:
        h_pos = "center"
    
    if center_y < 0.33:
        v_pos = "top"
    elif center_y > 0.66:
        v_pos = "bottom"
    else:
        v_pos = "middle"
    
    return f"{v_pos}-{h_pos}"


def _extract_layout_from_image(image_path: str, img_width: int, img_height: int) -> List[Dict[str, Any]]:
    """Extract text with bounding boxes from an image using Tesseract."""
    try:
        data = pytesseract.image_to_data(
            Image.open(image_path),
            output_type=pytesseract.Output.DICT,
            lang="eng"
        )
    except Exception as e:
        print(f"[DEBUG] Tesseract error: {e}")
        return []
    
    elements = []
    n_boxes = len(data["text"])
    
    for i in range(n_boxes):
        text = data["text"][i].strip()
        conf = data["conf"][i]
        
        if not text or conf < 0:
            continue
        
        bbox_raw = {
            "left": data["left"][i],
            "top": data["top"][i],
            "width": data["width"][i],
            "height": data["height"][i],
        }
        
        bbox = _normalize_bbox(bbox_raw, img_width, img_height)
        position = _get_position_label(bbox)
        
        elements.append({
            "index": len(elements),
            "text": text,
            "bbox": bbox,
            "confidence": round(conf / 100.0, 4),
            "position": position,
        })
    
    return elements


def _extract_pdf_with_layout(pdf_path: str) -> Dict[str, Any]:
    """Extract text with layout data from PDF."""
    print(f"[DEBUG] Starting PDF extraction: {pdf_path}")
    
    try:
        images = convert_from_path(pdf_path, dpi=200, poppler_path=POPPLER_PATH)
        print(f"[DEBUG] Converted PDF to {len(images)} page(s)")
    except Exception as e:
        raise ValueError(f"Failed to convert PDF to images: {e}")
    
    pages = []
    all_text_lines = []
    
    for page_num, image in enumerate(images):
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            image.save(tmp.name, "PNG")
            tmp_path = tmp.name
        
        try:
            img_width, img_height = image.size
            elements = _extract_layout_from_image(tmp_path, img_width, img_height)
            
            pages.append({
                "page": page_num + 1,
                "width": img_width,
                "height": img_height,
                "elements": elements,
            })
            
            for el in elements:
                all_text_lines.append(el["text"])
                
        finally:
            os.remove(tmp_path)
    
    layout_signature = _build_layout_signature(pages)
    
    return {
        "pages": pages,
        "full_text": " ".join(all_text_lines),
        "layout_signature": layout_signature,
        "total_elements": sum(len(p["elements"]) for p in pages),
    }


def _extract_docx_with_layout(docx_path: str) -> Dict[str, Any]:
    """Extract text from DOCX with basic structure."""
    doc = Document(docx_path)
    
    elements = []
    all_text = []
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        elements.append({
            "index": idx,
            "text": text,
            "type": "paragraph",
            "style": para.style.name if para.style else None,
        })
        all_text.append(text)
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    elements.append({
                        "index": len(elements),
                        "text": text,
                        "type": "table_cell",
                        "table": table_idx,
                        "row": row_idx,
                        "col": cell_idx,
                    })
                    all_text.append(text)
    
    return {
        "elements": elements,
        "full_text": "\n".join(all_text),
        "total_elements": len(elements),
        "source": "python-docx",
    }


def _build_layout_signature(pages: List[Dict]) -> Dict[str, Any]:
    """Build a layout signature for template comparison."""
    FIELD_PATTERNS = {
        "invoice_number": [
            "invoice no.", "invoice no", "invoice #", "invoice number", 
            "inv no", "inv #", "no:", "no.", "inv.", "invoice"
        ],
        "invoice_date": [
            "invoice date", "date:", "date", "issue date", "dated", 
            "01/", "02/", "03/", "04/", "05/", "06/", "07/", "08/", "09/", "10/", "11/", "12/"
        ],
        "due_date": ["due date", "payment due", "due:", "due "],
        "total": [
            "total", "grand total", "total amount", "amount due", 
            "balance due", "balance", "net amount", "total due"
        ],
        "subtotal": ["subtotal", "sub total", "sub-total", "sub total"],
        "tax": ["tax", "vat", "gst", "sales tax", "tax:"],
        "vendor": [
            "from:", "seller:", "vendor:", "bill from", "company", 
            "company name", "seller", "vendor"
        ],
        "customer": [
            "to:", "bill to:", "customer:", "client:", "sold to", 
            "issued to", "ship to", "deliver to", "issued", "billed to"
        ],
    }
    
    detected_fields = {}
    field_positions = {}
    
    if not pages:
        return {"fields": [], "positions": {}, "element_count": 0}
    
    first_page = pages[0]
    
    for element in first_page.get("elements", []):
        text_lower = element["text"].lower().strip()
        
        for field_name, patterns in FIELD_PATTERNS.items():
            if field_name in detected_fields:
                continue
                
            for pattern in patterns:
                if pattern in text_lower:
                    detected_fields[field_name] = {
                        "text": element["text"],
                        "bbox": element["bbox"],
                        "position": element["position"],
                        "confidence": element["confidence"],
                    }
                    field_positions[field_name] = element["position"]
                    break
    
    return {
        "fields": list(detected_fields.keys()),
        "positions": field_positions,
        "detected_fields": detected_fields,
        "element_count": len(first_page.get("elements", [])),
    }


def extract_text_with_layout(file_path: str, filename: str) -> Dict[str, Any]:
    """Full extraction with bounding boxes for template service."""
    name = filename.lower()
    
    if name.endswith(".pdf"):
        result = _extract_pdf_with_layout(file_path)
        result["source"] = "tesseract"
        return result
    
    if name.endswith(".docx"):
        return _extract_docx_with_layout(file_path)
    
    raise ValueError(f"Unsupported file type: {filename}")
