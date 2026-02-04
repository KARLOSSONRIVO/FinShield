"""
Layout Extraction (Template/AI).
"""
import os
import tempfile
from typing import Dict, Any, List
from pdf2image import convert_from_path
from docx import Document
import pytesseract
from PIL import Image
from app.utils.geometry import normalize_bbox, get_position_label
from ..signature import build_layout_signature


def _extract_layout_from_image(image_path: str, img_width: int, img_height: int) -> List[Dict[str, Any]]:
    """Extract text with bounding boxes from an image using Tesseract."""
    try:
        data = pytesseract.image_to_data(
            Image.open(image_path),
            output_type=pytesseract.Output.DICT,
            lang="eng"
        )
    except Exception as e:
        print(f"[DEBUG] Tesseract error: {e}")
        return []
    
    elements = []
    n_boxes = len(data["text"])
    
    for i in range(n_boxes):
        text = data["text"][i].strip()
        conf = data["conf"][i]
        
        if not text or conf < 0:
            continue
        
        bbox_raw = {
            "left": data["left"][i],
            "top": data["top"][i],
            "width": data["width"][i],
            "height": data["height"][i],
        }
        
        bbox = normalize_bbox(bbox_raw, img_width, img_height)
        position = get_position_label(bbox)
        
        elements.append({
            "index": len(elements),
            "text": text,
            "bbox": bbox,
            "confidence": round(conf / 100.0, 4),
            "position": position,
        })
    
    return elements


def _extract_pdf_with_layout(pdf_path: str) -> Dict[str, Any]:
    """Extract text with layout data from PDF."""
    print(f"[DEBUG] Starting PDF extraction: {pdf_path}")
    
    try:
        images = convert_from_path(pdf_path, dpi=200)
        print(f"[DEBUG] Converted PDF to {len(images)} page(s)")
    except Exception as e:
        raise ValueError(f"Failed to convert PDF to images: {e}")
    
    pages = []
    all_text_lines = []
    
    for page_num, image in enumerate(images):
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            image.save(tmp.name, "PNG")
            tmp_path = tmp.name
        
        try:
            img_width, img_height = image.size
            elements = _extract_layout_from_image(tmp_path, img_width, img_height)
            
            pages.append({
                "page": page_num + 1,
                "width": img_width,
                "height": img_height,
                "elements": elements,
            })
            
            for el in elements:
                all_text_lines.append(el["text"])
                
        finally:
            os.remove(tmp_path)
    
    layout_signature = build_layout_signature(pages)
    
    return {
        "pages": pages,
        "full_text": " ".join(all_text_lines),
        "layout_signature": layout_signature,
        "total_elements": sum(len(p["elements"]) for p in pages),
    }


def _extract_docx_with_layout(docx_path: str) -> Dict[str, Any]:
    """Extract text from DOCX with basic structure."""
    doc = Document(docx_path)
    
    # First pass: count total elements to normalize position
    temp_elements = []  # List of (text, type)
    all_text = []
    elements = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            temp_elements.append((text, "paragraph"))
            all_text.append(text)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    temp_elements.append((text, "table_cell"))
                    all_text.append(text)
                    
    total_count = len(temp_elements)
    
    # Second pass: build elements with simulated layout
    for i, (text, el_type) in enumerate(temp_elements):
        # Simulate Y position (linear from 0 to 1)
        y_pos = round(i / total_count, 4) if total_count > 0 else 0
        
        # Approximate geometric zones
        position = "middle-center"
        if y_pos < 0.33:
            position = "top-left"
        elif y_pos > 0.66:
            position = "bottom-left"
        else:
            position = "middle-left"

        # Simulate BBox
        bbox = {
            "x1": 0.1, 
            "y1": y_pos, 
            "x2": 0.9, 
            "y2": y_pos + (1/total_count if total_count else 0.05)
        }
        
        elements.append({
            "index": i,
            "text": text,
            "type": el_type,
            "bbox": bbox,
            "position": position,
            "confidence": 1.0, 
        })

    # Create a simulated "page" for the signature builder
    pages = [{
        "page": 1,
        "width": 1000,
        "height": 1000 * (total_count / 10 if total_count else 1),
        "elements": elements
    }]
    
    layout_signature = build_layout_signature(pages)
    
    return {
        "elements": elements,
        "pages": pages,
        "full_text": "\n".join(all_text),
        "total_elements": len(elements),
        "layout_signature": layout_signature,
        "source": "python-docx",
    }


def extract_text_with_layout(file_path: str, filename: str) -> Dict[str, Any]:
    """Full extraction with bounding boxes for template service."""
    name = filename.lower()
    
    if name.endswith(".pdf"):
        result = _extract_pdf_with_layout(file_path)
        result["source"] = "tesseract"
        return result
    
    if name.endswith(".docx"):
        return _extract_docx_with_layout(file_path)
    
    raise ValueError(f"Unsupported file type: {filename}")
