"""
Simple Text Extraction (OCR/Precheck).
"""
from docx import Document
import pdfplumber
import pytesseract


def extract_text_from_pdf(path: str) -> str:
    """Extract text from PDF using pdfplumber with table-aware extraction."""
    text_blocks = []

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            # First, try to extract tables
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    # Convert table to text preserving structure
                    for row in table:
                        # Join cells with spacing to preserve columns
                        row_text = "  ".join(str(cell) if cell else "" for cell in row)
                        if row_text.strip():
                            text_blocks.append(row_text)
            
            # Then extract regular text (outside tables)
            text = page.extract_text()
            if text:
                text_blocks.append(text)
            else:
                # Fallback to OCR for scanned pages
                img = page.to_image(resolution=300).original
                ocr_text = pytesseract.image_to_string(img, config="--psm 6")
                if ocr_text:
                    text_blocks.append(ocr_text)

    return "\n".join(text_blocks)


def extract_text_from_docx(path: str) -> str:
    """Extract text from DOCX."""
    doc = Document(path)
    blocks = []

    # Extract tables with preserved structure
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                # Join cells with spacing to preserve columns
                blocks.append("  ".join(row_text))

    # Extract paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text)

    extracted_text = "\n".join(blocks)
    
    # DEBUG: Print extracted text
    print("\n" + "="*60)
    print("🔍 DOCX EXTRACTION DEBUG")
    print("="*60)
    print(extracted_text)
    print("="*60 + "\n")
    
    return extracted_text


def extract_text_simple(path: str, filename: str) -> dict:
    """Simple text extraction for OCR and precheck services."""
    name = filename.lower()

    if name.endswith(".pdf"):
        return {
            "text": extract_text_from_pdf(path),
            "source": "tesseract"
        }

    if name.endswith(".docx"):
        return {
            "text": extract_text_from_docx(path),
            "source": "python-docx"
        }

    raise ValueError("UNSUPPORTED_FILE_TYPE")
