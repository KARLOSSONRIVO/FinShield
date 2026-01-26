from docx import Document

def extract_text_from_doc(path: str) -> str:
    doc = Document(path)
    blocks = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    blocks.append(cell.text)

    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text)

    return "\n".join(blocks)
